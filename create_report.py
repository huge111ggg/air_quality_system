from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT = Path("城市空气质量等级预测系统课程设计报告.docx")


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(12)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(12)


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("城市空气质量等级预测系统\n课程设计报告")
    run.bold = True
    run.font.name = "黑体"
    run.font.size = Pt(20)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("课程名称：《人工智能与大数据导论》\n姓名：__________  学号：__________  班级：__________\n完成日期：2026年6月")

    doc.add_heading("摘要", level=1)
    paragraph(
        doc,
        "本课程设计构建了一个城市空气质量等级预测系统。系统使用220条模拟空气质量监测数据，包含温度、湿度、风速、降雨量、交通拥堵指数、工业排放指数、PM2.5、PM10、NO2、SO2和臭氧等特征，将空气质量划分为优、良、轻度污染、中度及以上污染四类。项目完成了数据生成、数据清洗、异常值处理、特征标准化、模型训练、算法对比、可视化分析和Flask网页部署。实验对比了逻辑回归、随机森林和K近邻三种算法，最佳模型为逻辑回归，测试集准确率达到89.09%，Macro-F1达到84.40%。",
    )
    paragraph(doc, "关键词：空气质量；AQI；机器学习；大数据分析；Flask；Scikit-learn")

    doc.add_heading("一、项目背景与意义", level=1)
    paragraph(
        doc,
        "城市空气质量与居民健康、交通管理和环境治理密切相关。机动车排放、工业生产、道路扬尘和气象条件会共同影响AQI等级。机器学习方法能够从多维监测数据中学习污染物浓度、气象条件和人为活动指标与空气质量等级之间的关系，为环境预警和治理决策提供辅助支持。",
    )
    paragraph(
        doc,
        "本项目贴近环境监测这一社会热点问题，同时覆盖数据处理、特征工程、模型训练、模型评价和可视化部署等课程知识点，适合作为《人工智能与大数据导论》的综合实践项目。",
    )

    doc.add_heading("二、需求分析", level=1)
    paragraph(doc, "系统目标是根据气象指标、污染物浓度和人为活动指数预测空气质量等级，并展示模型训练结果。")
    bullet(doc, "数据需求：生成100至300条空气质量监测数据。")
    bullet(doc, "预处理需求：完成缺失值处理、异常值裁剪、特征标准化等处理。")
    bullet(doc, "建模需求：对比逻辑回归、随机森林和K近邻三种算法。")
    bullet(doc, "评价需求：输出准确率、Macro-F1、混淆矩阵和模型对比表。")
    bullet(doc, "展示需求：通过Flask网页展示运行结果、预测等级和治理建议。")

    doc.add_heading("三、系统设计", level=1)
    paragraph(
        doc,
        "系统采用数据层、模型层和应用层三层结构。数据层由generate_dataset.py生成模拟数据；模型层由train_model.py完成数据预处理、模型训练、算法对比和图表输出；应用层由app.py基于Flask实现网页预测功能。",
    )
    paragraph(
        doc,
        "系统流程为：生成监测数据，读取并清洗数据，对异常值进行裁剪，对数值特征进行标准化，划分训练集和测试集，训练三种分类模型，计算评价指标，保存最佳模型，最后在Flask网页中加载模型并进行在线预测。",
    )

    doc.add_heading("四、数据集与预处理", level=1)
    paragraph(
        doc,
        "数据集共220条记录，13个字段。输入特征包括温度、湿度、风速、降雨量、交通拥堵指数、工业排放指数、PM2.5、PM10、NO2、SO2和臭氧；aqi为空气质量指数，quality_level为空气质量等级标签。",
    )
    paragraph(
        doc,
        "数据预处理包括三类核心方法：第一，缺失值处理，使用中位数填充数值特征；第二，异常值处理，对温度、湿度、风速、降雨和污染物浓度进行合理范围裁剪；第三，特征标准化，使用StandardScaler统一不同指标的量纲。",
    )

    doc.add_heading("五、模型构建", level=1)
    paragraph(
        doc,
        "系统使用Scikit-learn构建机器学习流程，通过Pipeline和ColumnTransformer封装预处理与模型训练过程，保证训练阶段和预测阶段使用一致的数据处理逻辑。数据集按照75%训练集和25%测试集划分，并使用分层抽样保持类别比例相对稳定。",
    )
    paragraph(
        doc,
        "本项目对比逻辑回归、随机森林和K近邻三种分类算法。逻辑回归训练速度快、可解释性较强；随机森林能够学习非线性关系并提供特征重要性；K近邻可作为简单直观的基线模型。",
    )

    doc.add_heading("六、实验结果", level=1)
    paragraph(
        doc,
        "实验结果显示，逻辑回归准确率为89.09%，Macro-F1为84.40%；随机森林准确率为90.91%，Macro-F1为81.11%；K近邻准确率为81.82%，Macro-F1为57.28%。虽然随机森林准确率略高，但逻辑回归Macro-F1更好，说明其对不同类别的综合识别能力更均衡，因此系统选择逻辑回归作为最佳模型。",
    )
    paragraph(
        doc,
        "系统生成了模型性能对比图、混淆矩阵、空气质量等级分布图和关键因素图。这些图表能够展示模型性能、分类效果、数据分布以及影响AQI的关键因素。",
    )

    doc.add_heading("七、系统运行与展示", level=1)
    paragraph(
        doc,
        "在VS Code终端进入air_quality_system目录后，依次运行python generate_dataset.py、python train_model.py和python app.py。由于心理压力系统可能占用5000端口，本空气质量系统使用5001端口，浏览器访问http://127.0.0.1:5001即可打开网页。",
    )
    paragraph(
        doc,
        "用户输入温度、湿度、风速、降雨量、交通拥堵指数、工业排放指数和污染物浓度后，点击“预测空气质量”，系统会输出空气质量等级、各类别概率和治理建议。",
    )

    doc.add_heading("八、分析讨论", level=1)
    paragraph(
        doc,
        "逻辑回归和随机森林均取得较好结果，说明模拟数据中的污染物浓度与空气质量等级之间存在较清晰的关系。K近邻表现较弱，可能受到类别不均衡和多维特征距离度量的影响。",
    )
    paragraph(
        doc,
        "本项目的主要局限是数据为模拟生成，不能完全替代真实城市监测数据。后续可接入真实AQI开放数据，增加时间序列特征、季节特征和空间区域信息，并尝试XGBoost、LightGBM等模型。",
    )

    doc.add_heading("九、总结", level=1)
    paragraph(
        doc,
        "本课程设计完成了城市空气质量等级预测系统的设计与实现。项目实现了数据生成、数据预处理、三种机器学习模型对比、量化指标评价、可视化分析和Flask网页部署。系统能够根据输入监测指标预测空气质量等级并给出治理建议，体现了人工智能与大数据技术在环境监测领域的应用价值。",
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "李航. 统计学习方法[M]. 北京: 清华大学出版社, 2019.",
        "Ian Goodfellow, Yoshua Bengio, Aaron Courville. Deep Learning[M]. MIT Press, 2016.",
        "James, G., Witten, D., Hastie, T., Tibshirani, R. An Introduction to Statistical Learning[M]. Springer, 2021.",
        "Pedregosa, F. et al. Scikit-learn: Machine Learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "McKinney, W. Data Structures for Statistical Computing in Python[C]. Proceedings of the 9th Python in Science Conference, 2010.",
        "Hunter, J. D. Matplotlib: A 2D Graphics Environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.",
        "Breiman, L. Random Forests[J]. Machine Learning, 2001, 45: 5-32.",
        "中国环境监测总站. 空气质量指数AQI技术规定与监测方法相关资料[EB/OL].",
        "World Health Organization. WHO global air quality guidelines[EB/OL].",
    ]
    for ref in refs:
        paragraph(doc, ref)

    doc.add_heading("附录：运行命令", level=1)
    paragraph(doc, "cd C:\\codex工作空间\\skill-ppt\\air_quality_system")
    paragraph(doc, "python generate_dataset.py")
    paragraph(doc, "python train_model.py")
    paragraph(doc, "python app.py")
    paragraph(doc, "浏览器访问：http://127.0.0.1:5001")

    doc.save(OUTPUT)
    print(f"报告已生成：{OUTPUT.resolve()}")


if __name__ == "__main__":
    main()
